"""
import_docx_to_db.py - Extracts structured data from a Word .docx file
and stores it into the SQLite database using SQLAlchemy.
"""
from docx import Document
from sqlalchemy.exc import SQLAlchemyError
import os
from backend.database import engine, SessionLocal
from backend.models import Base, Question
from backend.config import SQLITE_DB_PATH

def init_db_if_needed():
    """创建数据库表如果它们不存在"""
    print(f"Checking database at {SQLITE_DB_PATH}")
    if not os.path.exists(SQLITE_DB_PATH) or not os.path.getsize(SQLITE_DB_PATH):
        print("Creating database tables...")
        Base.metadata.create_all(bind=engine)
        print("Database tables created.")
    else:
        print("Database already exists, ensuring tables...")
        # 确保所有表存在
        Base.metadata.create_all(bind=engine)

def load_questions_from_docx(docx_path):
    """从Word文档加载问题到数据库"""
    # 首先确保数据库和表存在
    init_db_if_needed()
    
    doc = Document(docx_path)
    session = SessionLocal()

    section = "Unknown Section"
    table_index = 0
    total_tables = len(doc.tables)
    imported_count = 0

    print("Importing from:", docx_path)

    for block in doc.element.body:
        if block.tag.endswith('}p'):
            text = ''.join(t.text for t in block.iter() if t.tag.endswith('}t')).strip()
            if text.lower().startswith("section"):
                section = text  
                print(f"Found section: {section}")

        elif block.tag.endswith('}tbl') and table_index < total_tables:
            table = doc.tables[table_index]
            table_index += 1

            for i, row in enumerate(table.rows):
                if i == 0:
                    continue  # skip header

                try:
                    seq = row.cells[0].text.strip()
                    page_name = row.cells[1].text.strip()
                    audio_file = row.cells[2].text.strip()
                    content = row.cells[3].text.strip()

                    if not content:
                        continue

                    q = Question(
                        section=section,
                        seq=seq,
                        page_name=page_name,
                        audio_file=audio_file,
                        content=content,
                        knowledge_point="To be added",  # 新增字段
                        question="To be added",         # 新增字段
                        answer="To be added"
                    )
                    session.add(q)
                    imported_count += 1
                    
                    # 每100条数据提交一次，减轻内存压力
                    if imported_count % 100 == 0:
                        try:
                            session.commit()
                            print(f"Committed {imported_count} records so far...")
                        except SQLAlchemyError as e:
                            print(f"Error in batch commit: {e}")
                            session.rollback()
                            
                except IndexError as e:
                    print(f"Error accessing table cells in row {i}, table {table_index}: {e}")
                except AttributeError as e:
                    print(f"Error reading cell content in row {i}, table {table_index}: {e}")
                except SQLAlchemyError as e:
                    print(f"Database error while adding row {i} from table {table_index}: {e}")
                    session.rollback()

    try:
        session.commit()
        print(f"Successfully imported {imported_count} records in total.")
    except SQLAlchemyError as e:
        print(f"Error committing changes to database: {e}")
        session.rollback()
    finally:
        session.close()
        print("Import complete.")

if __name__ == "__main__":
    load_questions_from_docx("data/Thunderstorm Avoidance_Boeing 20250210.docx")